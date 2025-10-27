"""
Enhanced document conversion engine supporting comprehensive document format conversions.
Uses existing libraries (python-docx, PyPDF2) with improved capabilities and error handling.
Optimized for performance with streaming, caching, and memory management.
"""

import os
import logging
import gc
import threading
from typing import Dict, List, Optional, Any, Generator
from pathlib import Path
import tempfile
import subprocess
from concurrent.futures import ThreadPoolExecutor
import time

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    from pdf2docx import Converter as PDFToDocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

from .base_engine import ConversionEngine, ConversionError

logger = logging.getLogger(__name__)

class DocumentEngine(ConversionEngine):
    """Enhanced document conversion engine with comprehensive format support and performance optimization."""
    
    def __init__(self):
        super().__init__("DocumentEngine")
        self._conversion_cache = {}
        self._cache_lock = threading.Lock()
        self._max_cache_size = 50  # Maximum cached conversions
        self._chunk_size = 8192  # 8KB chunks for streaming
        self._thread_pool = ThreadPoolExecutor(max_workers=2)  # Limit concurrent conversions
    
    def _initialize_formats(self) -> None:
        """Initialize supported document formats and conversion matrix."""
        self.supported_inputs = []
        self.supported_outputs = []
        
        # Add formats based on available libraries
        if PDF_AVAILABLE:
            self.supported_inputs.extend(['pdf'])
            self.supported_outputs.extend(['pdf'])
        
        if DOCX_AVAILABLE:
            self.supported_inputs.extend(['docx'])
            self.supported_outputs.extend(['docx', 'txt', 'html'])
        
        if XLSX_AVAILABLE:
            self.supported_inputs.extend(['xlsx'])
            self.supported_outputs.extend(['xlsx', 'csv'])
        
        # Text formats (always available)
        self.supported_inputs.extend(['txt', 'html', 'csv'])
        self.supported_outputs.extend(['txt', 'html', 'csv'])
        
        # Build conversion matrix
        self._build_conversion_matrix()
    
    def _build_conversion_matrix(self) -> None:
        """Build the conversion matrix based on available libraries."""
        # PDF conversions
        if PDF_AVAILABLE:
            self.conversion_matrix['pdf'] = ['txt']  # Basic text extraction
            if DOCX_AVAILABLE:
                self.conversion_matrix['pdf'].append('docx')
        
        # DOCX conversions
        if DOCX_AVAILABLE:
            self.conversion_matrix['docx'] = ['txt', 'html']
        
        # XLSX conversions
        if XLSX_AVAILABLE:
            self.conversion_matrix['xlsx'] = ['csv', 'txt', 'html']
            if PDF_AVAILABLE:
                self.conversion_matrix['xlsx'].append('pdf')
        
        # Text format conversions
        self.conversion_matrix['txt'] = ['html']
        if DOCX_AVAILABLE:
            self.conversion_matrix['txt'].append('docx')
        if REPORTLAB_AVAILABLE:
            self.conversion_matrix['txt'].append('pdf')
        
        self.conversion_matrix['html'] = ['txt']
        if DOCX_AVAILABLE:
            self.conversion_matrix['html'].append('docx')
        
        self.conversion_matrix['csv'] = ['txt', 'html']
        if XLSX_AVAILABLE:
            self.conversion_matrix['csv'].append('xlsx')
    
    def convert(self, input_path: str, output_path: str, 
                input_format: str, output_format: str, 
                options: Optional[Dict[str, Any]] = None) -> bool:
        """
        Convert between document formats with performance optimization.
        
        Options:
            - preserve_formatting: Boolean for format preservation
            - extract_images: Boolean for image extraction from documents
            - page_range: Tuple (start, end) for page-specific operations
            - use_cache: Boolean to enable/disable caching (default: True)
            - streaming: Boolean to enable streaming for large files (default: True)
        """
        start_time = time.time()
        try:
            options = options or {}
            input_format = input_format.lower()
            output_format = output_format.lower()
            
            if not self.can_convert(input_format, output_format):
                raise ConversionError(
                    f"Cannot convert {input_format} to {output_format}",
                    engine=self.name,
                    input_format=input_format,
                    output_format=output_format
                )
            
            # Check cache if enabled
            if options.get('use_cache', True):
                cache_key = self._get_cache_key(input_path, output_format, options)
                cached_result = self._get_from_cache(cache_key)
                if cached_result and os.path.exists(cached_result):
                    logger.info(f"Using cached conversion result: {cached_result}")
                    # Copy cached file to output path
                    import shutil
                    shutil.copy2(cached_result, output_path)
                    return True
            
            # Get file size for optimization decisions
            file_size = os.path.getsize(input_path)
            is_large_file = file_size > 5 * 1024 * 1024  # 5MB threshold
            
            # Route to appropriate conversion method
            conversion_key = f"{input_format}_to_{output_format}"
            
            # Add performance options
            perf_options = options.copy()
            perf_options.update({
                'is_large_file': is_large_file,
                'file_size': file_size,
                'streaming': options.get('streaming', True)
            })
            
            success = False
            if hasattr(self, f"_convert_{conversion_key}"):
                success = getattr(self, f"_convert_{conversion_key}")(input_path, output_path, perf_options)
            else:
                success = self._generic_convert(input_path, output_path, input_format, output_format, perf_options)
            
            # Cache successful conversion if enabled
            if success and options.get('use_cache', True):
                self._add_to_cache(cache_key, output_path)
            
            # Force garbage collection for large files
            if is_large_file:
                gc.collect()
            
            conversion_time = time.time() - start_time
            logger.info(f"Conversion completed in {conversion_time:.2f}s (file size: {file_size/1024/1024:.1f}MB)")
            
            return success
                
        except Exception as e:
            logger.error(f"Document conversion failed: {str(e)}")
            raise ConversionError(f"Document conversion failed: {str(e)}", engine=self.name)
    
    def _get_cache_key(self, input_path: str, output_format: str, options: Dict[str, Any]) -> str:
        """Generate a cache key for the conversion."""
        import hashlib
        
        # Get file modification time and size for cache invalidation
        stat = os.stat(input_path)
        file_info = f"{stat.st_mtime}_{stat.st_size}"
        
        # Create hash of input path, output format, and relevant options
        cache_data = f"{input_path}_{output_format}_{file_info}"
        
        # Include relevant options in cache key
        relevant_options = {k: v for k, v in options.items() 
                          if k in ['preserve_formatting', 'extract_images', 'page_range']}
        if relevant_options:
            cache_data += f"_{str(sorted(relevant_options.items()))}"
        
        return hashlib.md5(cache_data.encode()).hexdigest()
    
    def _get_from_cache(self, cache_key: str) -> Optional[str]:
        """Get cached conversion result."""
        with self._cache_lock:
            return self._conversion_cache.get(cache_key)
    
    def _add_to_cache(self, cache_key: str, output_path: str) -> None:
        """Add conversion result to cache."""
        with self._cache_lock:
            # Create cache directory if it doesn't exist
            cache_dir = os.path.join(tempfile.gettempdir(), 'docswap_cache')
            os.makedirs(cache_dir, exist_ok=True)
            
            # Copy file to cache with unique name
            cache_file = os.path.join(cache_dir, f"{cache_key}_{os.path.basename(output_path)}")
            import shutil
            shutil.copy2(output_path, cache_file)
            
            # Add to cache
            self._conversion_cache[cache_key] = cache_file
            
            # Limit cache size
            if len(self._conversion_cache) > self._max_cache_size:
                # Remove oldest entry
                oldest_key = next(iter(self._conversion_cache))
                old_file = self._conversion_cache.pop(oldest_key)
                if os.path.exists(old_file):
                    os.remove(old_file)
    
    def _convert_pdf_to_txt(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Extract text from PDF with optimized performance."""
        if not PDF_AVAILABLE:
            return False
        
        try:
            is_large_file = options.get('is_large_file', False)
            streaming = options.get('streaming', True)
            
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                page_range = options.get('page_range')
                if page_range:
                    start_page, end_page = page_range
                    pages = range(start_page - 1, min(end_page, total_pages))
                else:
                    pages = range(total_pages)
                
                # For large files or when streaming is enabled, process pages in chunks
                if is_large_file and streaming:
                    with open(output_path, 'w', encoding='utf-8') as output_file:
                        chunk_size = 10  # Process 10 pages at a time
                        for i in range(0, len(pages), chunk_size):
                            chunk_pages = pages[i:i + chunk_size]
                            chunk_text = []
                            
                            for page_num in chunk_pages:
                                try:
                                    page = pdf_reader.pages[page_num]
                                    text = page.extract_text()
                                    if text.strip():  # Only add non-empty text
                                        chunk_text.append(text)
                                except Exception as e:
                                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                                    continue
                            
                            if chunk_text:
                                output_file.write('\n\n'.join(chunk_text))
                                if i + chunk_size < len(pages):  # Not the last chunk
                                    output_file.write('\n\n')
                                output_file.flush()  # Ensure data is written
                            
                            # Force garbage collection for large files
                            if is_large_file:
                                gc.collect()
                else:
                    # Standard processing for smaller files
                    text_content = []
                    for page_num in pages:
                        try:
                            page = pdf_reader.pages[page_num]
                            text = page.extract_text()
                            if text.strip():  # Only add non-empty text
                                text_content.append(text)
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                            continue
                    
                    with open(output_path, 'w', encoding='utf-8') as output_file:
                        output_file.write('\n\n'.join(text_content))
                
                logger.info(f"Successfully extracted text from {len(pages)} pages")
                return True
                
        except Exception as e:
            logger.error(f"PDF to text conversion failed: {str(e)}")
            return False
    
    def _convert_pdf_to_docx(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert PDF to DOCX using pdf2docx with optimized performance."""
        if not PDF2DOCX_AVAILABLE:
            logger.error("pdf2docx library not available")
            return False
        
        try:
            is_large_file = options.get('is_large_file', False)
            file_size = options.get('file_size', 0)
            
            logger.info(f"Starting PDF to DOCX conversion (file size: {file_size/1024/1024:.1f}MB)")
            
            # Create converter instance with optimized settings
            cv = PDFToDocxConverter(input_path)
            
            # Get page range if specified
            page_range = options.get('page_range')
            
            # For large files, use optimized conversion parameters
            conversion_params = {}
            if is_large_file:
                # Optimize for large files
                conversion_params.update({
                    'multi_processing': True,  # Enable multiprocessing if available
                    'cpu_count': min(2, os.cpu_count() or 1),  # Limit CPU usage
                })
                logger.info("Using optimized settings for large file conversion")
            
            # Perform conversion with progress tracking
            start_time = time.time()
            
            if page_range:
                start_page, end_page = page_range
                cv.convert(output_path, start=start_page-1, end=end_page-1, **conversion_params)
                logger.info(f"Converted pages {start_page}-{end_page}")
            else:
                cv.convert(output_path, **conversion_params)
                logger.info("Converted all pages")
            
            cv.close()
            
            # Force garbage collection for large files
            if is_large_file:
                gc.collect()
            
            conversion_time = time.time() - start_time
            
            # Verify output file was created and has content
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                output_size = os.path.getsize(output_path)
                logger.info(f"Successfully converted PDF to DOCX in {conversion_time:.2f}s (output: {output_size/1024/1024:.1f}MB)")
                return True
            else:
                logger.error("PDF to DOCX conversion produced empty file")
                return False
                
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {str(e)}")
            # Clean up partial files
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
            return False
    
    def _convert_docx_to_txt(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Extract text from DOCX."""
        if not DOCX_AVAILABLE:
            return False
        
        try:
            doc = Document(input_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write('\n'.join(text_content))
            
            return True
            
        except Exception as e:
            logger.error(f"DOCX to text conversion failed: {str(e)}")
            return False
    
    def _convert_docx_to_html(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert DOCX to HTML."""
        if not DOCX_AVAILABLE:
            return False
        
        try:
            doc = Document(input_path)
            html_content = ['<!DOCTYPE html>', '<html>', '<head>', 
                          '<meta charset="utf-8">', '<title>Document</title>', 
                          '</head>', '<body>']
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    html_content.append(f'<p>{paragraph.text}</p>')
            
            html_content.extend(['</body>', '</html>'])
            
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write('\n'.join(html_content))
            
            return True
            
        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed: {str(e)}")
            return False
    
    def _convert_txt_to_docx(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert text to DOCX."""
        if not DOCX_AVAILABLE:
            return False
        
        try:
            doc = Document()
            
            with open(input_path, 'r', encoding='utf-8') as input_file:
                content = input_file.read()
            
            # Split into paragraphs
            paragraphs = content.split('\n\n')
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            logger.error(f"Text to DOCX conversion failed: {str(e)}")
            return False
    
    def _convert_txt_to_pdf(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert text to PDF."""
        if not REPORTLAB_AVAILABLE:
            return False
        
        try:
            # Read the text file
            with open(input_path, 'r', encoding='utf-8') as input_file:
                content = input_file.read()
            
            # Create PDF
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Set up text formatting
            margin = 72  # 1 inch margin
            line_height = 14
            max_width = width - 2 * margin
            y_position = height - margin
            
            # Split content into lines
            lines = content.split('\n')
            
            for line in lines:
                # Handle long lines by wrapping
                if c.stringWidth(line) > max_width:
                    words = line.split(' ')
                    current_line = ''
                    
                    for word in words:
                        test_line = current_line + (' ' if current_line else '') + word
                        if c.stringWidth(test_line) <= max_width:
                            current_line = test_line
                        else:
                            if current_line:
                                c.drawString(margin, y_position, current_line)
                                y_position -= line_height
                                if y_position < margin:
                                    c.showPage()
                                    y_position = height - margin
                            current_line = word
                    
                    if current_line:
                        c.drawString(margin, y_position, current_line)
                        y_position -= line_height
                else:
                    c.drawString(margin, y_position, line)
                    y_position -= line_height
                
                # Check if we need a new page
                if y_position < margin:
                    c.showPage()
                    y_position = height - margin
            
            c.save()
            return True
            
        except Exception as e:
            logger.error(f"Text to PDF conversion failed: {str(e)}")
            return False
    
    def _convert_txt_to_html(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert text to HTML."""
        try:
            # Read the text file
            with open(input_path, 'r', encoding='utf-8') as input_file:
                content = input_file.read()
            
            # Create HTML structure
            html_content = [
                '<!DOCTYPE html>',
                '<html lang="en">',
                '<head>',
                '    <meta charset="UTF-8">',
                '    <meta name="viewport" content="width=device-width, initial-scale=1.0">',
                '    <title>Converted Document</title>',
                '    <style>',
                '        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }',
                '        .content { max-width: 800px; margin: 0 auto; }',
                '        p { margin-bottom: 1em; }',
                '        pre { background-color: #f4f4f4; padding: 10px; border-radius: 4px; overflow-x: auto; }',
                '    </style>',
                '</head>',
                '<body>',
                '    <div class="content">'
            ]
            
            # Process content - split into paragraphs and handle line breaks
            paragraphs = content.split('\n\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Handle single line breaks within paragraphs
                    lines = paragraph.strip().split('\n')
                    if len(lines) == 1:
                        # Single line paragraph
                        html_content.append(f'        <p>{self._escape_html(lines[0])}</p>')
                    else:
                        # Multi-line paragraph - preserve line breaks
                        html_content.append('        <p>')
                        for i, line in enumerate(lines):
                            if i > 0:
                                html_content.append('            <br>')
                            html_content.append(f'            {self._escape_html(line)}')
                        html_content.append('        </p>')
            
            # Close HTML structure
            html_content.extend([
                '    </div>',
                '</body>',
                '</html>'
            ])
            
            # Write HTML file
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write('\n'.join(html_content))
            
            logger.info("Successfully converted text to HTML")
            return True
            
        except Exception as e:
            logger.error(f"Text to HTML conversion failed: {str(e)}")
            return False
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#x27;'))
    
    def _convert_xlsx_to_csv(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert XLSX to CSV."""
        if not XLSX_AVAILABLE:
            return False
        
        try:
            workbook = openpyxl.load_workbook(input_path)
            worksheet = workbook.active
            
            import csv
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                csv_writer = csv.writer(csvfile)
                
                for row in worksheet.iter_rows(values_only=True):
                    csv_writer.writerow(row)
            
            return True
            
        except Exception as e:
            logger.error(f"XLSX to CSV conversion failed: {str(e)}")
            return False
    
    def _convert_csv_to_xlsx(self, input_path: str, output_path: str, options: Dict[str, Any]) -> bool:
        """Convert CSV to XLSX."""
        if not XLSX_AVAILABLE:
            return False
        
        try:
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            
            import csv
            with open(input_path, 'r', encoding='utf-8') as csvfile:
                csv_reader = csv.reader(csvfile)
                
                for row_num, row in enumerate(csv_reader, 1):
                    for col_num, value in enumerate(row, 1):
                        worksheet.cell(row=row_num, column=col_num, value=value)
            
            workbook.save(output_path)
            return True
            
        except Exception as e:
            logger.error(f"CSV to XLSX conversion failed: {str(e)}")
            return False
    
    def _generic_convert(self, input_path: str, output_path: str, 
                        input_format: str, output_format: str, options: Dict[str, Any]) -> bool:
        """Generic conversion fallback."""
        # For now, return False for unsupported conversions
        # This can be enhanced with additional conversion logic
        logger.warning(f"Generic conversion not implemented for {input_format} to {output_format}")
        return False
    
    def get_available_features(self) -> Dict[str, bool]:
        """Get information about available features."""
        return {
            'pdf_support': PDF_AVAILABLE,
            'docx_support': DOCX_AVAILABLE,
            'xlsx_support': XLSX_AVAILABLE,
            'text_extraction': True,
            'html_generation': True,
            'csv_support': True
        }